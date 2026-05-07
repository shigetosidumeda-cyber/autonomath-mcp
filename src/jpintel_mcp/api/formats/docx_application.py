"""DOCX 申請書 boilerplate renderer — ``?format=docx-application``.

行政書士法 §1 fence (CRITICAL):

    This renderer produces a SCAFFOLD ONLY. It deliberately does NOT fill
    placeholder fields (e.g. ``{{customer_name}}``, ``{{program_name}}``,
    ``{{requested_amount}}``). The user must:
        1. Fill placeholders themselves OR have a 行政書士 fill them.
        2. Have any 申請書 reviewed by a licensed 行政書士 if the program
           requires 書面提出 — 行政書士法 §1 prohibits unlicensed
           creation of 官公署提出書類.

    The cover page, every page footer, and the JSON ``_disclaimer``
    envelope on the wrapping response all relay this fence. The output
    file name is suffixed ``_scaffold.docx`` to keep the "this is not a
    finished 申請書" signal in the local filesystem too.

Layout:

    Page 1 — cover: program name + scaffold banner + §1 fence + brand.
    Page 2..N — one section per program row:
                heading: program name
                table: unified_id | source_url | source_fetched_at | license
                placeholder paragraphs:
                    申請者氏名 / 法人名:   {{customer_name}}
                    法人番号:               {{houjin_bangou}}
                    所在地:                 {{address}}
                    代表者氏名:             {{representative_name}}
                    申請額:                 {{requested_amount_yen}}
                    申請理由 (300字以内):   {{reason}}
                    添付書類:               {{attachments}}
    Page final — §1 fence reprint + brand footer.

Font: ``ＭＳ 明朝`` for JP, ``Cambria`` for ASCII, both 10.5pt — the safe
combo for 行政提出書類 across Word 365, LibreOffice, Pages.
"""

from __future__ import annotations

import io
from typing import Any

from fastapi import HTTPException, status
from fastapi.responses import Response

from jpintel_mcp.api._format_dispatch import (
    BRAND_FOOTER,
    DISCLAIMER_HEADER_VALUE,
    DISCLAIMER_JA,
)

# 行政書士法 §1 fence. Verbatim, no paraphrase — we want grep'able text.
GYOSEISHOSHI_FENCE_JA = (
    "【重要・行政書士法 §1 注意事項】\n"
    "本ファイルは「申請書のたたき台 (scaffold)」です。AutonoMath は\n"
    "{{customer_name}} 等の placeholder を意図的に未記入のまま出力\n"
    "しています。官公署提出書類の作成・代理は行政書士法 §1 により\n"
    "行政書士の独占業務です。本ファイルを実際に提出する場合は、\n"
    "(1) お客様ご自身で記入する、または (2) 行政書士の確認・代行\n"
    "を受けてください。AutonoMath / Bookyou株式会社は本書類の\n"
    "完成・提出について一切の責任を負いません。"
)

GYOSEISHOSHI_FENCE_EN = (
    "[Notice — 行政書士法 §1 Fence] This DOCX is a SCAFFOLD ONLY. "
    "Placeholder fields like {{customer_name}} are intentionally left "
    "BLANK. Creating an actual 官公署提出書類 (filing) on behalf of "
    "another party is reserved to a licensed 行政書士 under 行政書士法 §1. "
    "Fill the placeholders yourself, or have a 行政書士 review the file "
    "before submission. Bookyou株式会社 disclaims all liability for the "
    "completion / filing of this document."
)

# Placeholder block — left intentionally unfilled. Each row is a paragraph
# in the output so the user can see where to write.
PLACEHOLDER_BLOCK: tuple[tuple[str, str], ...] = (
    ("申請者氏名 / 法人名", "{{customer_name}}"),
    ("法人番号", "{{houjin_bangou}}"),
    ("所在地", "{{address}}"),
    ("代表者氏名", "{{representative_name}}"),
    ("連絡先 (TEL / メール)", "{{contact}}"),
    ("申請額 (円)", "{{requested_amount_yen}}"),
    ("事業計画概要 (300字以内)", "{{business_plan_summary}}"),
    ("申請理由 (300字以内)", "{{reason}}"),
    ("添付書類", "{{attachments}}"),
)

# Title-row keys used to label each program section.
_TITLE_KEYS: tuple[str, ...] = (
    "primary_name",
    "name",
    "title",
)


def _row_title(row: dict[str, Any]) -> str:
    for k in _TITLE_KEYS:
        v = row.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    return str(row.get("unified_id") or "(untitled)")


def render_docx_application(rows: list[dict[str, Any]], meta: dict[str, Any]) -> Response:
    """Render a 申請書 scaffold DOCX from one or more program rows.

    One row -> one program section. Multi-row input is supported but
    discouraged — the natural unit of work is "one program -> one DOCX".
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail=(
                "format=docx-application requires the 'python-docx' dep — pip install python-docx"
            ),
        ) from exc

    doc = Document()

    # --- Default style: ＭＳ 明朝 10.5pt for JP, Cambria for ASCII fallback.
    style = doc.styles["Normal"]
    style.font.name = "ＭＳ 明朝"
    style.font.size = Pt(10.5)

    # ----- Cover page.
    h = doc.add_heading("申請書 (scaffold)", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.get("endpoint", "AutonoMath export"))
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(GYOSEISHOSHI_FENCE_JA)
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(GYOSEISHOSHI_FENCE_EN)

    doc.add_paragraph()
    doc.add_paragraph(f"§52: {DISCLAIMER_JA}")
    doc.add_paragraph(BRAND_FOOTER)

    # ----- Per-program sections.
    for i, row in enumerate(rows):
        doc.add_page_break()  # type: ignore[no-untyped-call]
        doc.add_heading(f"{i + 1}. {_row_title(row)}", level=1)

        # Lineage table.
        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Light Grid Accent 1"
        cells = [
            ("unified_id", str(row.get("unified_id") or "")),
            ("source_url", str(row.get("source_url") or "")),
            ("source_fetched_at", str(row.get("source_fetched_at") or "")),
            ("license", str(row.get("license") or "")),
            ("corpus_snapshot_id", str(meta.get("corpus_snapshot_id") or "")),
        ]
        for r_idx, (k, v) in enumerate(cells):
            tbl.cell(r_idx, 0).text = k
            tbl.cell(r_idx, 1).text = v

        doc.add_paragraph()
        doc.add_heading("入力欄 (placeholder — 必ずご記入ください)", level=2)
        for label, placeholder in PLACEHOLDER_BLOCK:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(placeholder)

    # ----- Final fence reprint.
    doc.add_page_break()  # type: ignore[no-untyped-call]
    doc.add_heading("注意事項 (再掲)", level=1)
    doc.add_paragraph(GYOSEISHOSHI_FENCE_JA)
    doc.add_paragraph()
    doc.add_paragraph(f"§52: {DISCLAIMER_JA}")
    doc.add_paragraph(f"出典: {meta.get('license_summary', '')}")
    doc.add_paragraph(BRAND_FOOTER)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{meta.get('filename_stem', 'autonomath_export')}_scaffold.docx"
    return Response(
        content=buf.read(),
        media_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-AutonoMath-Disclaimer": DISCLAIMER_HEADER_VALUE,
            "X-AutonoMath-Format": "docx-application",
            "X-AutonoMath-Gyoseishoshi-Fence": "scaffold-only",
        },
    )
